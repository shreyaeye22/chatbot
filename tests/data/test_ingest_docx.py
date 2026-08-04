from __future__ import annotations

import pytest
from docx import Document

from data.db import get_connection, init_db
from data.ingest_docx import extract_text_from_docx, ingest_docx


@pytest.fixture
def sample_docx_path(tmp_path):
    path = tmp_path / "worksheet.docx"
    doc = Document()
    doc.add_paragraph("Comparative and superlative worksheet")
    doc.add_paragraph("Complete exercises 1 to 10.")
    doc.save(path)
    return str(path)


def test_extract_text_from_docx_returns_paragraph_text(sample_docx_path):
    text = extract_text_from_docx(sample_docx_path)

    assert "Comparative and superlative worksheet" in text
    assert "Complete exercises 1 to 10." in text


def test_ingest_docx_stores_a_course_content_row(tmp_path, sample_docx_path):
    db_path = tmp_path / "app.db"
    init_db(db_path)
    conn = get_connection(db_path)

    row_id = ingest_docx(
        conn,
        file=sample_docx_path,
        subject="math",
        topic="comparatives",
        source_name="worksheet.docx",
    )

    row = conn.execute(
        "SELECT subject, topic, content, source FROM course_content WHERE id = ?", (row_id,)
    ).fetchone()
    conn.close()

    assert row["subject"] == "math"
    assert row["topic"] == "comparatives"
    assert "Complete exercises 1 to 10." in row["content"]
    assert row["source"] == "worksheet.docx"


def test_ingest_docx_rejects_empty_document(tmp_path):
    empty_path = tmp_path / "empty.docx"
    Document().save(empty_path)

    db_path = tmp_path / "app.db"
    init_db(db_path)
    conn = get_connection(db_path)

    with pytest.raises(ValueError):
        ingest_docx(
            conn, file=str(empty_path), subject="math", topic="x", source_name="empty.docx"
        )
    conn.close()
